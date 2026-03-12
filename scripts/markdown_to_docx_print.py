#!/usr/bin/env python3
"""Convert a constrained Markdown document into a print-ready DOCX."""

from __future__ import annotations

import argparse
import re
from pathlib import Path
from typing import List, Tuple

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


LINK_RE = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")
INLINE_RE = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`)")
HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")
ORDERED_RE = re.compile(r"^\s*(\d+)\.\s+(.*)$")
UNORDERED_RE = re.compile(r"^\s*[-*]\s+(.*)$")
TABLE_SEP_RE = re.compile(r"^\s*\|?(?:\s*:?-{3,}:?\s*\|)+\s*:?-{3,}:?\s*\|?\s*$")


def _clean_inline(text: str) -> str:
    return LINK_RE.sub(lambda m: f"{m.group(1)} ({m.group(2)})", text.strip())


def _set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def _shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def _set_cell_margins(cell, top=80, start=90, bottom=80, end=90) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_1)
    run._r.append(instr)
    run._r.append(fld_char_2)


def _parse_table(lines: List[str], start: int) -> Tuple[int, List[List[str]]]:
    table_lines = []
    i = start
    while i < len(lines):
        raw = lines[i]
        if not raw.strip():
            break
        if "|" not in raw:
            break
        table_lines.append(raw.rstrip())
        i += 1
    if len(table_lines) < 2 or not TABLE_SEP_RE.match(table_lines[1]):
        return start, []
    rows = []
    for idx, line in enumerate(table_lines):
        if idx == 1:
            continue
        stripped = line.strip().strip("|")
        rows.append([cell.strip() for cell in stripped.split("|")])
    return i, rows


def _add_inline_runs(paragraph, text: str) -> None:
    text = _clean_inline(text)
    pos = 0
    for match in INLINE_RE.finditer(text):
        if match.start() > pos:
            paragraph.add_run(text[pos:match.start()])
        token = match.group(0)
        if token.startswith("**") and token.endswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("`") and token.endswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)
        pos = match.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def _add_multiline_runs(paragraph, lines: List[Tuple[str, bool]]) -> None:
    force_break_all = len(lines) > 1 and all(text.startswith("**") for text, _ in lines)
    for idx, (text, hard_break) in enumerate(lines):
        if idx > 0:
            if force_break_all or lines[idx - 1][1]:
                paragraph.add_run().add_break()
            else:
                paragraph.add_run(" ")
        _add_inline_runs(paragraph, text)


def _configure_document(doc: Document, title: str) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.0)
    section.header_distance = Cm(0.9)
    section.footer_distance = Cm(1.0)
    section.different_first_page_header_footer = True

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)

    for style_name, size, color in (
        ("Title", 22, RGBColor(20, 49, 83)),
        ("Heading 1", 16, RGBColor(20, 49, 83)),
        ("Heading 2", 13, RGBColor(39, 71, 122)),
        ("Heading 3", 11.5, RGBColor(39, 71, 122)),
    ):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color

    header = section.header.paragraphs[0]
    header.text = title
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if header.runs:
        header.runs[0].font.name = "Calibri"
        header.runs[0].font.size = Pt(9)
        header.runs[0].font.color.rgb = RGBColor(110, 110, 110)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Página ")
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    _add_page_number(footer)


def _add_table(doc: Document, rows: List[List[str]]) -> None:
    if not rows:
        return
    col_count = max(len(row) for row in rows)
    table = doc.add_table(rows=1, cols=col_count)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = True

    header_cells = table.rows[0].cells
    for idx in range(col_count):
        cell = header_cells[idx]
        _set_cell_margins(cell)
        _shade_cell(cell, "143153")
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        text = rows[0][idx] if idx < len(rows[0]) else ""
        _add_inline_runs(para, text)
        for run in para.runs:
            run.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.size = Pt(10)
    _set_repeat_table_header(table.rows[0])

    for row_idx, row_data in enumerate(rows[1:], start=1):
        row = table.add_row()
        for col_idx in range(col_count):
            cell = row.cells[col_idx]
            _set_cell_margins(cell)
            if row_idx % 2 == 0:
                _shade_cell(cell, "F5F7FA")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            text = row_data[col_idx] if col_idx < len(row_data) else ""
            _add_inline_runs(para, text)
            for run in para.runs:
                run.font.size = Pt(9.5)
    doc.add_paragraph()


def convert_markdown_to_docx(source: Path, dest: Path) -> None:
    lines = source.read_text(encoding="utf-8").splitlines()
    doc = Document()
    title = source.stem
    _configure_document(doc, title)

    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        if stripped in {"---", "***"}:
            i += 1
            continue

        table_end, rows = _parse_table(lines, i)
        if rows:
            _add_table(doc, rows)
            i = table_end
            continue

        heading = HEADING_RE.match(stripped)
        if heading:
            level = min(len(heading.group(1)), 3)
            text = heading.group(2).strip()
            if text.lower().startswith("anexo "):
                doc.add_page_break()
            para = doc.add_paragraph(style=f"Heading {level}" if level > 1 else "Title")
            if level == 1:
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            _add_inline_runs(para, text)
            if level == 1:
                doc.add_paragraph()
            i += 1
            continue

        unordered = UNORDERED_RE.match(line)
        if unordered:
            while i < len(lines):
                match = UNORDERED_RE.match(lines[i])
                if not match:
                    break
                para = doc.add_paragraph(style="List Bullet")
                para.paragraph_format.space_after = Pt(1)
                _add_inline_runs(para, match.group(1).strip())
                i += 1
            continue

        ordered = ORDERED_RE.match(line)
        if ordered:
            while i < len(lines):
                match = ORDERED_RE.match(lines[i])
                if not match:
                    break
                para = doc.add_paragraph(style="List Number")
                para.paragraph_format.space_after = Pt(1)
                _add_inline_runs(para, match.group(2).strip())
                i += 1
            continue

        paragraph_lines = [(stripped, line.endswith("  "))]
        i += 1
        while i < len(lines):
            raw_next = lines[i].rstrip("\n")
            nxt = raw_next.strip()
            if not nxt:
                break
            if HEADING_RE.match(nxt) or UNORDERED_RE.match(lines[i]) or ORDERED_RE.match(lines[i]):
                break
            table_end, rows = _parse_table(lines, i)
            if rows:
                break
            if nxt in {"---", "***"}:
                break
            paragraph_lines.append((nxt, raw_next.endswith("  ")))
            i += 1
        para = doc.add_paragraph()
        para.paragraph_format.space_after = Pt(6)
        para.paragraph_format.line_spacing = 1.12
        _add_multiline_runs(para, paragraph_lines)

    dest.parent.mkdir(parents=True, exist_ok=True)
    doc.save(dest)


def main() -> None:
    parser = argparse.ArgumentParser(description="Convert markdown to print-ready DOCX.")
    parser.add_argument("source", type=Path)
    parser.add_argument("dest", type=Path)
    args = parser.parse_args()
    convert_markdown_to_docx(args.source, args.dest)


if __name__ == "__main__":
    main()
